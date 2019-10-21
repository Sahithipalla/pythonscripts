import docx
d=docx.Document("test.docx")
d.paragraphs
d.paragraphs[0]
d.paragraphs[0].text
d.paragraphs[1].text
for p in d.paragraphs:
     print(p.text)
#     or
     
p1=d.paragraphs[0]
p2=d.paragraphs[1]
p1.text
p2.text

#runs

p1.runs
p1.runs[0].text
p1.runs[1].text
p1.runs[2].text
p1.runs[3].text
p1.runs[4].text
p1.runs[5].text
p1.runs[7].text
p1.runs[2].underline=True
p1.runs[2].text="c programming"
d.save("test2.docx")

p2.runs
p2.runs[0].text
p2.runs[1].text
p2.runs[2].text
p2.runs[3].text

#create a word doc
d=docx.Document()
d.add_heading("page1 heading",0)
d.add_paragraph("Hello there,good morning")
d.add_paragraph("I am working with word file")
d.save("test3.docx")
#add text in existing paragraph
d.paragraphs[1].text
p1=d.paragraphs[1]
p1.add_run(", Beautiful day ouside")
d.save("test4.docx")
#diff styles in word files
d.paragraphs[2].text
p2=d.paragraphs[2]
p2.add_run("I am bold")
p2.add_run("I am italic")
d.save("test5.docx")
#add a page break and start a new page
d.add_page_break()
d.add_heading("Heading 2",2)
d.save("test6.docx")
#add picture

d=docx.Document()
d.add_picture("img_3.png",width=docx.shared.Cm(6),height=docx.shared.Cm(6))
d.save("test7.docx")
